import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_assignment():
    doc = Document()

    # Page Setup: Letter (8.5 x 11 inches)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Set up styles
    styles = doc.styles

    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.line_spacing = 1.5

    # Main Heading (Heading 1)
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = None
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_style.paragraph_format.line_spacing = 1.5

    # Sub Heading (Heading 2)
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = None
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_style.paragraph_format.line_spacing = 1.5

    def add_heading(text, level=1):
        doc.add_paragraph(text, style=f'Heading {level}')

    def add_paragraph(text):
        doc.add_paragraph(text, style='Normal')

    def add_bullet(text):
        doc.add_paragraph(text, style='List Bullet')

    # A. Student Declaration
    add_heading('A. Student Declaration', 1)
    add_paragraph("I hereby declare that this assignment is my own work and effort. It has not been submitted anywhere for any award. Where other sources of information have been used, they have been acknowledged in the reference section.")
    doc.add_page_break()

    # B. Acknowledgment
    add_heading('B. Acknowledgment', 1)
    add_paragraph("I would like to express my sincere gratitude to my instructor for providing the necessary guidance and resources to complete this LAN design project. I also extend my appreciation to my peers for their constructive feedback during the research and development phases.")
    doc.add_page_break()

    # C. Table of Content
    add_heading('C. Table of Content', 1)
    add_paragraph("A. Student Declaration")
    add_paragraph("B. Acknowledgment")
    add_paragraph("C. Table of Content")
    add_paragraph("D. Table of Figure")
    add_paragraph("E. Executive Summary")
    add_paragraph("F. Full Assignment Tasks")
    add_paragraph("G. Conclusion")
    add_paragraph("H. References")
    doc.add_page_break()

    # D. Table of Figure
    add_heading('D. Table of Figure', 1)
    add_paragraph("Figure 1: Logical Network Topology Diagram")
    add_paragraph("Figure 2: Physical Floor Plan Layout")
    doc.add_page_break()

    # E. Executive Summary
    add_heading('E. Executive Summary', 1)
    add_paragraph("This report presents a comprehensive Local Area Network (LAN) design for TechSavvy Inc., a rapidly expanding tech startup. The company's relocation to a new three-story building, encompassing 30,000 square feet and accommodating 150 employees with plans to scale to 300, necessitates a robust, secure, and scalable infrastructure. The proposed network design integrates high-speed wired and wireless connectivity, ensuring low latency for critical applications such as software development and video conferencing. A hierarchical network topology featuring core, distribution, and access layers guarantees efficient data flow, fault tolerance, and seamless integration of Future technologies. Security is addressed through strict segmentation via VLANs, firewall implementation, and robust wireless encryption. This design ensures TechSavvy Inc. maintains a competitive edge through uninterrupted and secure digital operations.")
    doc.add_page_break()

    # F. Full Assignment Tasks
    add_heading('F. Full Assignment Tasks', 1)

    # Task 1
    add_heading('Task 1 - Project Requirements Analysis', 2)
    add_paragraph("a) Key Requirements Identification")
    add_paragraph("The primary requirement is to design a LAN supporting 150 current employees across three floors (10,000 sq. ft. each), with an infrastructure capable of supporting 300 users. The physical layout includes open-plan workspaces, 15 private offices, 4 video-conferencing rooms, a server room (second floor), and common areas. Departmental needs dictate diverse device support, including high-performance workstations for developers, laptops for management, VoIP phones, network printers, and wireless mobile devices. Specific needs include:")
    add_paragraph("- Development: Low latency, high bandwidth for large code deployments.")
    add_paragraph("- Management/HR/Finance: Secure access to sensitive data, segmented from the general network.")
    add_paragraph("- General: Reliable Wi-Fi coverage and seamless video conferencing capabilities.")
    
    add_paragraph("b) Growth and Scalability Analysis")
    add_paragraph("TechSavvy Inc.'s projected growth to 300 employees demands a highly scalable architecture. High-bandwidth applications, specifically 4K video conferencing, VoIP, and large file transfers typical in software development, generate significant network traffic. This impacts the design by necessitating high-throughput backbone connections (e.g., 10Gbps fiber links between floors) and high-density wireless access points. Utilizing a modular core and distribution layer approach allows for seamless port capacity upgrades without disrupting existing services, ensuring the network can absorb the doubled user base and evolving technological demands.")

    # Task 2
    add_heading('Task 2 - Network Topology Design', 2)
    add_paragraph("a) Floor Plan Layout Justification")
    add_paragraph("The physical layout strategically centralizes the server room on the second floor, minimizing cable runs to the first and third floors to remain well within the 100-meter Ethernet standard. IDF (Intermediate Distribution Frame) closets on the first and third floors house access switches, connecting back to the MDF (Main Distribution Frame) in the second-floor server room. Access points are distributed in a honeycomb pattern across open-plan areas to provide overlapping coverage without interference, eliminating dead zones in conference and common rooms. (Refer to Figure 2 for the theoretical layout diagram).")
    
    add_paragraph("b) Logical Network Topology")
    add_paragraph("A collapsed core topology is recommended. This combines the core and distribution layers into high-capacity multilayer switches in the server room, providing routing and high-speed switching. Access layer switches on each floor connect endpoint devices. This model reduces latency, provides high redundancy via EtherChannel connections, and optimizes data flow. VLANs separate traffic logically (e.g., VLAN 10 for Development, VLAN 20 for Management).")
    p = doc.add_paragraph("[Insert Logical Network Topology Diagram Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Figure 1: Logical Network Topology Diagram")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Task 3
    add_heading('Task 3 - Equipment Selection and Specification', 2)
    add_paragraph("a) Required Networking Equipment")
    add_paragraph("- Routers: 2x Enterprise-grade Edge Routers (e.g., Cisco ISR 4000 series) for redundant WAN connectivity.")
    add_paragraph("- Switches (Core/Distribution): 2x Layer 3 Managed Switches (e.g., Cisco Catalyst 9300 series) for high-speed routing and redundancy.")
    add_paragraph("- Switches (Access): 6x 48-port PoE+ Layer 2 Switches (e.g., Cisco Catalyst 9200L) distributed across floors.")
    add_paragraph("- Wireless Access Points: 15x Wi-Fi 6 (802.11ax) Access Points (e.g., Cisco Meraki MR46) for high-density coverage.")
    add_paragraph("- Cables: Cat6a UTP for gigabit endpoint connections; OM4 Multimode Fiber for 10Gbps uplinks between floors.")
    
    add_paragraph("b) Equipment Justification")
    add_paragraph("Cisco enterprise equipment ensures high reliability, extensive support, and advanced security features. Cat6a cabling future-proofs endpoint connections up to 10Gbps, exceeding current bandwidth needs. PoE+ switches eliminate the need for separate power supplies for VoIP phones and Access Points, simplifying deployment. Wi-Fi 6 APs handle the dense deployment of wireless devices in open workspaces efficiently, providing the required throughput for high-bandwidth tasks.")

    # Task 4
    add_heading('Task 4 - IP Addressing and Subnetting', 2)
    add_paragraph("a) IP Addressing Plan")
    add_paragraph("A private Class B address space, specifically 172.16.0.0/16, provides an expansive foundation for routing and management. Using Classless Inter-Domain Routing (CIDR), the network will be subnetted using /24 masks for typical departments, allowing 254 usable hosts per subnet, optimizing broadcast domains.")
    
    add_paragraph("b) Subnet Allocation")
    add_paragraph("- VLAN 10 (Development): 172.16.10.0/24")
    add_paragraph("- VLAN 20 (Management/HR): 172.16.20.0/24")
    add_paragraph("- VLAN 30 (Servers/Data Center): 172.16.30.0/24")
    add_paragraph("- VLAN 40 (VoIP): 172.16.40.0/24")
    add_paragraph("- VLAN 50 (Wi-Fi Corporate): 172.16.50.0/24")
    add_paragraph("- VLAN 60 (Wi-Fi Guest): 172.16.60.0/24")
    
    add_paragraph("c) Future Growth Considerations")
    add_paragraph("The 172.16.0.0/16 block provides 256 /24 subnets. With only 6 initially deployed, there is immense scalability to accommodate the projected growth to 300 employees and any unforeseen departmental subdivisions without requiring IP readdressing.")

    # Task 5
    add_heading('Task 5 - Network Security Design', 2)
    add_paragraph("a) Security Threats and Vulnerabilities")
    add_paragraph("TechSavvy faces threats including unauthorized internal access to proprietary software code, malware/ransomware infections introduced via remote workers or phishing, and external DoS/DDoS attacks disrupting internet connectivity. Open-plan wireless access introduces risks of rogue access points or eavesdropping.")
    
    add_paragraph("b) Security Plan and Controls")
    add_paragraph("- Next-Generation Firewall (NGFW): Deployed at the network edge to perform deep packet inspection, intrusion prevention (IPS), and content filtering.")
    add_paragraph("- Network Access Control (NAC): IEEE 802.1X authentication for both wired and wireless connections ensures only authorized corporate devices can access internal VLANs.")
    add_paragraph("- WPA3 Enterprise: Implemented on all corporate wireless networks for robust encryption.")
    add_paragraph("- Segmentation: VLANs strictly separate guest traffic from corporate data, and inter-VLAN routing is heavily controlled by Access Control Lists (ACLs) on the core switches.")
    add_paragraph("Justification: These layered controls provide defense-in-depth, mitigating both internal and external threats while protecting intellectual property.")

    # Task 6
    add_heading('Task 6 - Network Management and Monitoring', 2)
    add_paragraph("a) Monitoring Tools")
    add_paragraph("A comprehensive monitoring solution such as SolarWinds Network Performance Monitor (NPM) or PRTG Network Monitor will be utilized. These tools leverage SNMP (Simple Network Management Protocol) and NetFlow to monitor bandwidth utilization, device health, and network latency in real-time.")
    
    add_paragraph("b) Reliability and Uptime")
    add_paragraph("Reliability is ensured through redundant hardware, including dual power supplies in core switches and dual internet service providers (ISPs) utilizing HSRP (Hot Standby Router Protocol) for seamless failover. The monitoring software will be configured with automated alerting thresholds (e.g., CPU utilization > 80% or link down), allowing the IT support team to proactively address issues before they impact end-users, ensuring maximum uptime.")

    doc.add_page_break()

    # G. Conclusion
    add_heading('G. Conclusion', 1)
    add_paragraph("The network architecture designed for TechSavvy Inc. delivers a resilient, high-performance, and secure infrastructure tailored to their dynamic startup environment. By leveraging a collapsed core topology, enterprise-grade hardware, and a scalable IP addressing scheme, the LAN will effortlessly support both current operations and the anticipated growth to 300 employees. Strict security measures and proactive monitoring ensure that TechSavvy's proprietary data remains protected while maintaining continuous operational uptime.")
    doc.add_page_break()

    # H. References
    add_heading('H. References', 1)
    add_paragraph("Cisco Systems (2026) 'Campus LAN and Wireless LAN Solution Design Guide'. Available at: https://www.cisco.com (Accessed: 2 July 2026).")
    add_paragraph("Stallings, W. (2021) 'Data and Computer Communications', 10th edn. Pearson Education.")
    add_paragraph("Kurose, J.F. and Ross, K.W. (2021) 'Computer Networking: A Top-Down Approach', 8th edn. Pearson.")

    # Save document
    save_path = r'c:\Users\User\Desktop\NEW\TechSavvy_Network_Design_Assignment.docx'
    doc.save(save_path)
    print(f"Assignment generated successfully at: {save_path}")

if __name__ == '__main__':
    create_assignment()
