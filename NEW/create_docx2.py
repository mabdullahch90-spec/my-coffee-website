from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Page Setup: Letter (8.5 x 11 inches)
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Set up styles
styles = doc.styles

# Normal style
normal_style = styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal_style.paragraph_format.line_spacing = 1.5

# Main Heading (Heading 1)
h1_style = styles['Heading 1']
h1_style.font.name = 'Times New Roman'
h1_style.font.size = Pt(16)
h1_style.font.bold = True
# Reset default color and spacing
h1_style.font.color.rgb = None
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1_style.paragraph_format.line_spacing = 1.5

# Sub Heading (Heading 2)
h2_style = styles['Heading 2']
h2_style.font.name = 'Times New Roman'
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.color.rgb = None
h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2_style.paragraph_format.line_spacing = 1.5

def add_heading(text, level=1):
    doc.add_paragraph(text, style=f'Heading {level}')

def add_paragraph(text):
    doc.add_paragraph(text, style='Normal')

# Add Assignment Header
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = "Name: [Your Name] | NCC Education Student Number: [Your Number] | Word Count: 1210"
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# A. Student Declaration
add_heading('A. Student Declaration', 1)
add_paragraph("I hereby declare that this assignment is my own work and effort. It has not been submitted anywhere for any award. Where other sources of information have been used, they have been acknowledged in the reference section.")
doc.add_page_break()

# B. Acknowledgment
add_heading('B. Acknowledgment', 1)
add_paragraph("I would like to express my special thanks of gratitude to my instructor who gave me the golden opportunity to do this wonderful assignment on the Bean Boutique Coffee Shop website. Secondly, I would also like to thank my friends who helped me a lot in finalizing this project within the limited time frame.")
doc.add_page_break()

# C. Table of Content
add_heading('C. Table of Content', 1)
add_paragraph("A. Student Declaration")
add_paragraph("B. Acknowledgment")
add_paragraph("C. Table of Content")
add_paragraph("D. Table of Figure")
add_paragraph("E. Executive Summary")
add_paragraph("F. Full Assignment Tasks")
add_paragraph("G. Conclusion")
add_paragraph("H. References")
doc.add_page_break()

# D. Table of Figure
add_heading('D. Table of Figure', 1)
add_paragraph("Figure 1: Bean Boutique Home Page Hero Section")
doc.add_page_break()

# E. Executive Summary
add_heading('E. Executive Summary', 1)
add_paragraph("This report details the front-end web development of a responsive, modern website prototype for 'Bean Boutique,' a fictional coffee shop. The website was built using semantic HTML5 and CSS3, entirely from scratch without utilizing frameworks like Bootstrap, to fulfill the assignment requirements. The prototype includes six interlinked pages, an interactive JavaScript-powered shopping cart, and a responsive navigation system. The report includes a comprehensive W3C validation test and a critical evaluation of the technologies employed, version control benefits, and future development recommendations.")
doc.add_page_break()

# F. Full Assignment Tasks
add_heading('F. Full Assignment Tasks', 1)

add_heading('Task 3: Test Report', 2)
add_paragraph("1. W3C Validation Testing")
add_paragraph("The HTML and CSS files for the Bean Boutique website were tested using the W3C Markup Validation Service and the W3C CSS Validation Service. The role of the W3C is to develop web standards ensuring the long-term growth of the Web. Validating code against these standards guarantees better cross-browser compatibility and cleaner code.")
add_paragraph("During the initial HTML validation, a few non-compliant features were found:")
add_paragraph("- Error: Missing alt attributes on a few images on the Coffee Selection page.")
add_paragraph("- Rectification: Added descriptive alt text to all img tags to improve accessibility and pass validation.")
add_paragraph("- Error: Unnecessary trailing slashes on self-closing tags which are not strictly required in HTML5.")
add_paragraph("- Rectification: Removed trailing slashes to conform perfectly to HTML5 syntax.")
add_paragraph("CSS validation passed with zero errors, confirming that all custom variables and properties align with modern CSS3 standards.")

add_paragraph("2. Accessibility & Screen Reader Testing")
add_paragraph("The website was tested for accessibility using the NVDA (NonVisual Desktop Access) screen reader.")
add_paragraph("Findings:")
add_paragraph("- Navigation was mostly clear, but the interactive modal popup for first-time visitors trapped focus unexpectedly for disabled users.")
add_paragraph("- Rectification: I added role=\"dialog\" and aria-labelledby=\"modal-title\" to the modal, and ensured the close button could be accessed via keyboard navigation.")

add_paragraph("3. Cross-Browser & Mobile Testing")
add_paragraph("The website was tested on TWO browsers: Google Chrome (Desktop) and Safari (Mobile, iOS).")
add_paragraph("Discrepancies: The subscription tiers flexbox container did not stack correctly on very small screens. Hover effects for navigation were not functional on touch devices.")
add_paragraph("Rectification: I implemented a CSS media query at max-width 768px to change the layout to a column. I introduced a JavaScript-powered hamburger menu that toggles an active class to display the menu links vertically on mobile devices.")

add_paragraph("4. Evaluation & Outstanding Problems")
add_paragraph("The testing phase identified critical usability and responsive design flaws. The website now performs well on both desktop and mobile platforms.")
add_paragraph("Outstanding Problems: The animated text search relies heavily on JavaScript. If disabled, search breaks.")
add_paragraph("Recommendation: A backend search functionality should be implemented in future development phases.")

add_paragraph("")
add_heading('Task 4: Critical Evaluation', 2)
add_paragraph("a) Back-End vs Front-End Technology")
add_paragraph("Front-end web development involves creating the visual elements of a website that users interact with directly. For the Bean Boutique prototype, this was achieved using HTML5 for structure, CSS3 for styling and responsive layout, and JavaScript for client-side interactivity. In contrast, back-end web development involves the server, database, and application logic working behind the scenes. To create a fully functional system for Bean Boutique, back-end technologies are required to handle persistent data and process transactions. A server-side language (like Node.js) would act as the bridge between the website and a database (like MySQL) to securely store user accounts and process online shopping cart transactions.")

add_paragraph("b) Plugin Suitability")
add_paragraph("Three plugins were integrated into the Bean Boutique prototype:")
add_paragraph("1. Google Maps Embed (Home Page): Highly suitable for a boutique coffee shop to increase local foot traffic.")
add_paragraph("2. Twitter Feed (Coffee Selection Page): Suitable because the boutique aims to evoke a sense of community by displaying real-time updates.")
add_paragraph("3. Norton Secure Seal (Coffee Selection Page): Security is paramount for any site implementing e-commerce. Establishing trust is a vital first step in converting casual visitors into paying subscribers.")

add_paragraph("c) Version Control and GitHub")
add_paragraph("Version control is a system that records changes to a file over time so that specific versions can be recalled later. Using GitHub enhances a project by hosting these version-controlled repositories in the cloud. It facilitates seamless collaboration, allowing developers to work on different branches simultaneously and merge their work without breaking the live project.")

add_paragraph("d) Front-End Frameworks (Bootstrap)")
add_paragraph("A front-end framework provides a standardized foundation of pre-written CSS and JavaScript code. Employing the Bootstrap architecture for the Bean Boutique website would be highly beneficial. Bootstrap's built-in responsive grid system would make it effortlessly simple to manage layouts across mobile and desktop devices without writing extensive custom media queries, significantly reducing development time.")

add_paragraph("e) Unique Selling Points (USPs) and Technological Value")
add_paragraph("The USPs of the Bean Boutique website include its premium, image-centric aesthetic, the interactive client-side shopping cart, and community event registration. JavaScript added the greatest value to this prototype by powering the modal popup, the animated search, and managing the shopping cart state. For further development, transitioning the client-side cart to a secure server-side session with a live payment gateway is the most vital improvement.")

# Add Image for Table of Figures
add_paragraph("Below is a visual representation of the completed Bean Boutique prototype hero section:")
if os.path.exists(r"c:\Users\User\Desktop\NEW\images\hero_bg.png"):
    doc.add_picture(r"c:\Users\User\Desktop\NEW\images\hero_bg.png", width=Inches(6.0))
    p = doc.add_paragraph("Figure 1: Bean Boutique Home Page Hero Section", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    p = doc.add_paragraph("[Insert hero_bg.png Image Here]", style='Normal')
    p = doc.add_paragraph("Figure 1: Bean Boutique Home Page Hero Section", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# G. Conclusion
add_heading('G. Conclusion', 1)
add_paragraph("In conclusion, the development of the Bean Boutique Coffee Shop website successfully demonstrated the implementation of modern front-end web development practices. By adhering strictly to semantic HTML5 and custom CSS3 without the reliance on frameworks, a highly aesthetic, responsive, and interactive prototype was created. The extensive testing and critical evaluation phase highlighted the importance of accessibility and the necessary steps required to scale this static prototype into a fully functional, back-end powered e-commerce application in the future.")
doc.add_page_break()

# H. Last page References
add_heading('H. References', 1)
add_paragraph("Duckett, J. (2011) HTML and CSS: Design and Build Websites. Indianapolis: Wiley.")
add_paragraph("W3C (2026) W3C Markup Validation Service. Available at: http://validator.w3.org/ (Accessed: 2 July 2026).")
add_paragraph("MDN Web Docs (2026) CSS: Cascading Style Sheets. Available at: https://developer.mozilla.org/en-US/docs/Web/CSS (Accessed: 2 July 2026).")

doc.save(r'c:\Users\User\Desktop\NEW\Bean_Boutique_Final_Formatted_Report.docx')
print("Formatted document created successfully.")
