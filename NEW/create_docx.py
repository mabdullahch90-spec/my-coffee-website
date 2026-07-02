from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Combine the text
test_report = """Task 3: Test Report

1. W3C Validation Testing
The HTML and CSS files for the Bean Boutique website were tested using the W3C Markup Validation Service (http://validator.w3.org/) and the W3C CSS Validation Service. The role of the W3C is to develop web standards ensuring the long-term growth of the Web. Validating code against these standards guarantees better cross-browser compatibility and cleaner code.

During the initial HTML validation, a few non-compliant features were found:
- Error: Missing alt attributes on a few images on the Coffee Selection page.
- Rectification: Added descriptive alt text to all img tags to improve accessibility and pass validation.
- Error: Unnecessary trailing slashes on self-closing tags which are not strictly required in HTML5.
- Rectification: Removed trailing slashes to conform perfectly to HTML5 syntax.

CSS validation passed with zero errors, confirming that all custom variables and properties align with modern CSS3 standards.

2. Accessibility & Screen Reader Testing
The website was tested for accessibility using the NVDA (NonVisual Desktop Access) screen reader.
Findings:
- Navigation was mostly clear, but the interactive modal popup for first-time visitors trapped focus unexpectedly for disabled users.
- Rectification: I added role="dialog" and aria-labelledby="modal-title" to the modal, and ensured the close button could be accessed via keyboard navigation.
- Form inputs were initially missing linked labels. I added label tags (sometimes hidden visually via a CSS class) to ensure screen readers announce input fields correctly, such as the newsletter email field and event registration form.

3. Cross-Browser & Mobile Testing
The website was tested on TWO browsers:
1. Google Chrome (Desktop, Windows 11): The site rendered perfectly. Hover effects, box shadows, and the animated text search function performed smoothly.
2. Safari (Mobile, iOS/iPhone):
- Discrepancies: The subscription tiers flexbox container did not stack correctly on very small screens, causing horizontal scrolling. Additionally, hover effects for the navigation menu are not functional on touch devices.
- Rectification: I implemented a CSS media query at max-width 768px to change the subscription tiers layout from a row layout to a column layout. For navigation, I introduced a JavaScript-powered hamburger menu that toggles an active class to display the menu links vertically on mobile devices, circumventing the reliance on hover states.

4. Evaluation & Outstanding Problems
The testing phase was suitable and effective for a static prototype, identifying critical usability and responsive design flaws. The website now performs well on both desktop and mobile platforms.

Outstanding Problems:
- The animated text search on the Coffee Selection page relies heavily on JavaScript. If a user has JavaScript disabled, the search functionality breaks.
- Recommendation: A backend search functionality should be implemented in future development phases. This would ensure the search works server-side regardless of client-side script settings, providing a more robust fallback."""

critical_eval = """Task 4: Critical Evaluation

a) Back-End vs Front-End Technology
Front-end web development involves creating the visual elements of a website that users interact with directly. For the Bean Boutique prototype, this was achieved using HTML5 for structure, CSS3 for styling and responsive layout, and JavaScript for client-side interactivity (such as the cart calculation and modal popup). Front-end development focuses on the user interface and user experience within the web browser.

In contrast, back-end web development involves the server, database, and application logic working behind the scenes. To create a fully functional system for Bean Boutique, back-end technologies are required to handle persistent data and process transactions. A server-side language (like Node.js, Python, or PHP) would act as the bridge between the website and a database (like MySQL or MongoDB). The database would securely store user accounts, product inventory, and order history. Furthermore, integrating a payment gateway API (like Stripe or PayPal) on the back-end would be essential to securely process the online shopping cart transactions and manage the proposed coffee subscription models.

b) Plugin Suitability
Three plugins were integrated into the Bean Boutique prototype to enhance functionality without requiring complex custom code:
1. Google Maps Embed (Home Page): This interactive map widget is highly suitable for a boutique coffee shop that operates a physical neighbourhood store. It allows users to find the exact location, pan, zoom, and get directions directly from the website. I chose this because location visibility is critical for local retail foot traffic.
2. Twitter/Social Media Feed (Coffee Selection Page): A social media block was included to highlight the shop's community engagement. This plugin is suitable because the boutique aims to evoke a sense of community. Displaying real-time updates about fresh roasts or events keeps the website content dynamic and encourages visitors to follow the brand's social channels.
3. Norton Secure Seal / Trust Badge (Coffee Selection Page): Security is paramount for any site planning to implement e-commerce. Including a security trust badge reassures customers that their data (like emails and future payment details) will be protected. I chose this specific visual cue because establishing trust is a vital first step in converting casual visitors into paying subscribers.

c) Version Control and GitHub
Version control is a system that records changes to a file or set of files over time so that specific versions can be recalled later. It allows developers to track modifications, revert to previous states if mistakes occur, and branch off to experiment with new features without breaking the main live project.

Using GitHub enhances a project by hosting these version-controlled repositories in the cloud. For Bean Boutique, GitHub would act as a secure backup for the source code. Furthermore, it facilitates seamless collaboration; if the boutique hires additional developers or designers, they can simultaneously work on different branches (e.g., one working on the cart functionality, another on UI tweaks) and merge their work via Pull Requests. GitHub also integrates with modern web hosting platforms for continuous deployment, meaning pushing a change to the main branch can automatically update the live website.

d) Front-End Frameworks (Bootstrap)
A front-end development framework is a pre-prepared software framework designed to support the development of web applications. It provides a standardized foundation of pre-written CSS and JavaScript code (like grid systems, typography, and UI components) which developers can build upon. The primary advantage is speed; developers do not have to write basic structural CSS from scratch, ensuring rapid prototyping and consistent cross-browser compatibility.

Employing the Bootstrap architecture for the Bean Boutique website would be highly beneficial. Bootstrap's built-in responsive grid system would make it effortlessly simple to manage the catalogue layouts across mobile, tablet, and desktop devices without writing extensive custom media queries. Additionally, its pre-styled components (like navigation bars, modals, and buttons) would significantly reduce development time, allowing the developer to focus more on custom branding overrides rather than foundational structural bugs.

e) Unique Selling Points (USPs) and Technological Value
The Unique Selling Points of the Bean Boutique website include its premium, image-centric aesthetic that minimizes text clutter, the interactive client-side shopping cart prototype, and the integrated community event registration.

The technology that added the greatest value to this prototype is JavaScript. While HTML and CSS established the visual foundation, JavaScript breathed life into the site. It powered the modal popup to capture first-time visitor emails, enabled the animated text search for quick product discovery, and most importantly, managed the shopping cart state using LocalStorage.

For further development, I recommend implementing a robust back-end framework to transition the prototype into a fully functioning e-commerce platform. Specifically, integrating a headless CMS (Content Management System) would allow the shop owners to easily update coffee blends and events without touching the HTML code. Finally, transitioning the client-side cart to a secure server-side session with a live payment gateway is the most vital improvement to realise the business's online revenue goals."""

# Combine text and count words
full_text = test_report + "\n\n" + critical_eval
# Basic word count splitting by whitespace
word_count = len(full_text.split())

doc = Document()

# Setup header
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = f"Name: [Your Name Here]   |   NCC ID: [Your ID Here]   |   Word Count: {word_count} / 1250"
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Apply document-wide default styling
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing = 2.0

for line in full_text.split('\n'):
    if line.strip() != "":
        p = doc.add_paragraph(line)

doc.save(r'c:\Users\User\Desktop\NEW\Bean_Boutique_Assignment_Report.docx')
print("Document created successfully.")
